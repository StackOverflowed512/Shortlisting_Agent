import PyPDF2 # Keep PyPDF2, as it's generally lighter if it works for your PDFs
from docx import Document
import logging
import os
from typing import Optional # Import Optional

logger = logging.getLogger(__name__)

def extract_text_from_pdf(file_path: str) -> Optional[str]: # Changed here
    try:
        with open(file_path, 'rb') as file:
            # Use PdfReader for PyPDF2 v3.0.0+
            # For older versions (like 1.x, 2.x), it was PdfFileReader
            try:
                reader = PyPDF2.PdfReader(file) # For PyPDF2 3.0.0+
            except AttributeError:
                # Fallback for older PyPDF2 versions if PdfReader is not found
                reader = PyPDF2.PdfFileReader(file) # type: ignore

            text = ""
            if hasattr(reader, 'pages'): # PyPDF2 3.0.0+
                for page_num in range(len(reader.pages)):
                    page = reader.pages[page_num]
                    text += page.extract_text() or "" # Ensure None is handled
            elif hasattr(reader, 'getNumPages'): # Older PyPDF2
                for page_num in range(reader.getNumPages()):
                    page = reader.getPage(page_num)
                    text += page.extractText() or "" # Ensure None is handled
            else:
                logger.error(f"Unsupported PyPDF2 version or invalid PDF object for {file_path}")
                return None

            if not text.strip():
                 logger.warning(f"No text extracted from PDF (possibly image-based or empty): {os.path.basename(file_path)}")
            else:
                logger.info(f"Successfully extracted text from PDF: {os.path.basename(file_path)}")
            return text
    except FileNotFoundError:
        logger.error(f"PDF file not found: {file_path}")
        return None
    except PyPDF2.errors.PdfReadError as e: # More specific error for PyPDF2
        logger.error(f"Error reading PDF (possibly corrupted or password-protected) {file_path}: {e}")
        return None
    except Exception as e:
        logger.error(f"Error extracting text from PDF {file_path}: {e}", exc_info=True)
        return None

def extract_text_from_docx(file_path: str) -> Optional[str]: # Changed here
    try:
        doc = Document(file_path)
        text = "\n".join([paragraph.text for paragraph in doc.paragraphs])
        if not text.strip():
            logger.warning(f"No text extracted from DOCX (possibly empty): {os.path.basename(file_path)}")
        else:
            logger.info(f"Successfully extracted text from DOCX: {os.path.basename(file_path)}")
        return text
    except FileNotFoundError:
        logger.error(f"DOCX file not found: {file_path}")
        return None
    except Exception as e: # Catches errors from python-docx, e.g., if file is not a valid DOCX
        logger.error(f"Error extracting text from DOCX {file_path}: {e}", exc_info=True)
        return None

def parse_resume(file_path: str) -> Optional[str]: # Changed here
    _, extension = os.path.splitext(file_path)
    extension = extension.lower()

    if extension == '.pdf':
        return extract_text_from_pdf(file_path)
    elif extension == '.docx':
        return extract_text_from_docx(file_path)
    else:
        logger.warning(f"Unsupported file type for parsing: {file_path}. Only PDF and DOCX are supported.")
        return None

if __name__ == '__main__':
    logging.basicConfig(level=logging.INFO, format='%(asctime)s - %(levelname)s - %(message)s')
    # Create dummy files for testing
    dummy_dir = "dummy_test_files"
    if not os.path.exists(dummy_dir):
        os.makedirs(dummy_dir)

    # --- Create a more realistic dummy PDF with text ---
    try:
        from reportlab.pdfgen import canvas
        from reportlab.lib.pagesizes import letter

        dummy_pdf_path = os.path.join(dummy_dir, "test.pdf")
        c = canvas.Canvas(dummy_pdf_path, pagesize=letter)
        c.drawString(72, 720, "This is a test PDF document.")
        c.drawString(72, 700, "It contains some sample text for PDF parsing.")
        c.drawString(72, 680, "Contact: test@example.com")
        c.save()
        print(f"Created dummy PDF: {dummy_pdf_path}")
    except ImportError:
        print("reportlab not installed. Skipping detailed dummy PDF creation. PDF parsing test might be limited.")
        # Fallback to simple blank PDF if reportlab is not available
        dummy_pdf_path = os.path.join(dummy_dir, "test_blank.pdf")
        if not os.path.exists(dummy_pdf_path): # Only create if reportlab failed
            try:
                from PyPDF2 import PdfWriter
                writer = PdfWriter()
                writer.add_blank_page(width=letter[0], height=letter[1])
                with open(dummy_pdf_path, "wb") as f_blank_pdf:
                    writer.write(f_blank_pdf)
                print(f"Created blank dummy PDF: {dummy_pdf_path}")
            except Exception as e_pypdf_write:
                print(f"Could not create even a blank PDF: {e_pypdf_write}")
                dummy_pdf_path = None


    dummy_docx_path = os.path.join(dummy_dir, "test.docx")
    doc = Document()
    doc.add_paragraph("This is a test DOCX file.")
    doc.add_paragraph("It contains some sample text for parsing.")
    doc.add_paragraph("Candidate: John Doe, Email: john.doe@example.net")
    doc.save(dummy_docx_path)
    print(f"Created dummy DOCX: {dummy_docx_path}")


    print("\nTesting PDF parsing:")
    if dummy_pdf_path and os.path.exists(dummy_pdf_path):
        pdf_text = parse_resume(dummy_pdf_path)
        if pdf_text:
            print(f"Extracted PDF Text:\n'{pdf_text[:200]}...'")
        else:
            print("Could not parse PDF or it was empty.")
    else:
        print("Dummy PDF not created, skipping PDF parsing test.")

    print("\nTesting DOCX parsing:")
    if os.path.exists(dummy_docx_path):
        docx_text = parse_resume(dummy_docx_path)
        if docx_text:
            print(f"Extracted DOCX Text:\n'{docx_text[:200]}...'")
        else:
            print("Could not parse DOCX.")
    else:
        print("Dummy DOCX not created, skipping DOCX parsing test.")

    # print("\nConsider cleaning up dummy files manually from 'dummy_test_files' directory if needed.")
    print("\nTo test with actual content, place sample PDF/DOCX files and update paths or run main.py.")